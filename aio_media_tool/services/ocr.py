from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from threading import Event

from aio_media_tool.models import JobCancelled
from aio_media_tool.services.common import (
    ProgressCallback,
    atomic_write_bytes,
    check_cancelled,
    noop_progress,
    unique_output,
)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"}


@dataclass(frozen=True, slots=True)
class OCRResult:
    source: str
    engine: str
    pages: tuple[str, ...]

    @property
    def text(self) -> str:
        if len(self.pages) <= 1:
            return self.pages[0] if self.pages else ""
        return "\n\n".join(
            f"--- Seite {index} ---\n{page}" for index, page in enumerate(self.pages, 1)
        )


class OCRService:
    @staticmethod
    def availability(tesseract_path: str = "") -> dict[str, str]:
        tesseract = tesseract_path.strip() or shutil.which("tesseract") or ""
        try:
            import easyocr  # noqa: F401

            easy = "verfügbar"
        except (ImportError, OSError):
            easy = "nicht installiert"
        try:
            import pymupdf  # noqa: F401

            pdf = "verfügbar"
        except (ImportError, OSError):
            pdf = "PyMuPDF fehlt"
        return {
            "Tesseract": tesseract or "nicht gefunden",
            "EasyOCR": easy,
            "PDF-Rendering": pdf,
        }

    def recognize(
        self,
        source: Path,
        engine: str = "auto",
        languages: str = "deu+eng",
        tesseract_path: str = "",
        private_temp: Path | None = None,
        progress: ProgressCallback = noop_progress,
        cancel: Event | None = None,
    ) -> OCRResult:
        if not source.is_file():
            raise FileNotFoundError(source)
        engine = engine.casefold()
        executable = tesseract_path.strip() or shutil.which("tesseract") or ""
        if engine == "auto":
            engine = "tesseract" if executable else "easyocr"
        if engine == "tesseract" and not executable:
            raise RuntimeError(
                "Tesseract-OCR wurde nicht gefunden. Pfad auswählen oder installieren."
            )
        if source.suffix.casefold() == ".pdf":
            images = self._render_pdf(source)
        elif source.suffix.casefold() in IMAGE_EXTENSIONS:
            images = [source]
        else:
            raise ValueError("OCR unterstützt PDF, JPG, PNG, TIFF, BMP und WebP.")
        pages: list[str] = []
        temp_parent = private_temp or source.parent
        temp_parent.mkdir(parents=True, exist_ok=True)
        with tempfile.TemporaryDirectory(prefix="aio-ocr-", dir=temp_parent) as directory:
            for index, image_source in enumerate(images):
                check_cancelled(cancel)
                if isinstance(image_source, bytes):
                    image_path = Path(directory) / f"page-{index + 1:05d}.png"
                    image_path.write_bytes(image_source)
                else:
                    image_path = image_source
                if engine == "tesseract":
                    text = self._tesseract(executable, image_path, languages, cancel)
                elif engine == "easyocr":
                    text = self._easyocr(image_path, languages)
                else:
                    raise ValueError(f"Unbekannte OCR-Engine: {engine}")
                pages.append(text.strip())
                progress(
                    int((index + 1) / len(images) * 100),
                    f"OCR-Seite {index + 1} von {len(images)} erkannt",
                )
        return OCRResult(str(source), engine, tuple(pages))

    @staticmethod
    def _render_pdf(source: Path) -> list[bytes]:
        try:
            import pymupdf
        except (ImportError, OSError) as exc:
            raise RuntimeError(
                "Für OCR auf PDF-Dateien fehlt PyMuPDF. Installiere 'uv sync --extra ocr'."
            ) from exc
        pages: list[bytes] = []
        with pymupdf.open(source) as document:
            for page in document:
                pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2.2, 2.2), alpha=False)
                pages.append(pixmap.tobytes("png"))
        if not pages:
            raise ValueError("Das PDF enthält keine Seiten.")
        return pages

    @staticmethod
    def _tesseract(executable: str, image_path: Path, languages: str, cancel: Event | None) -> str:
        process = subprocess.Popen(
            [
                executable,
                str(image_path),
                "stdout",
                "-l",
                languages or "eng",
                "--psm",
                "3",
                "-c",
                "preserve_interword_spaces=1",
            ],
            stdin=subprocess.DEVNULL,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            encoding="utf-8",
            errors="replace",
            shell=False,
        )
        try:
            while True:
                check_cancelled(cancel)
                try:
                    stdout, stderr = process.communicate(timeout=0.2)
                    break
                except subprocess.TimeoutExpired:
                    continue
        except JobCancelled:
            process.terminate()
            try:
                process.wait(timeout=3)
            except subprocess.TimeoutExpired:
                process.kill()
            raise
        if process.returncode:
            raise RuntimeError(stderr.strip()[-1500:] or "Tesseract-OCR ist fehlgeschlagen.")
        return stdout

    @staticmethod
    def _easyocr(image_path: Path, languages: str) -> str:
        try:
            import easyocr
        except (ImportError, OSError) as exc:
            raise RuntimeError("EasyOCR ist nicht installiert.") from exc
        mapping = {"deu": "de", "eng": "en", "fra": "fr", "spa": "es", "ita": "it"}
        selected = [mapping.get(code.casefold(), code.casefold()) for code in languages.split("+")]
        selected = [value for value in selected if value] or ["en"]
        try:
            reader = easyocr.Reader(selected, gpu=True, verbose=False)
        except (RuntimeError, OSError):
            reader = easyocr.Reader(selected, gpu=False, verbose=False)
        raw = reader.readtext(str(image_path), detail=1, paragraph=False)
        ordered = sorted(
            raw,
            key=lambda item: (
                min(point[1] for point in item[0]),
                min(point[0] for point in item[0]),
            ),
        )
        return "\n".join(str(item[1]) for item in ordered)

    @staticmethod
    def translate_deepl(
        text: str,
        api_key: str,
        target_language: str,
        use_free_api: bool = True,
        progress: ProgressCallback = noop_progress,
        cancel: Event | None = None,
    ) -> str:
        if not api_key.strip():
            raise ValueError("Bitte einen DeepL-API-Key eingeben.")
        if not text.strip():
            return ""
        chunks = OCRService._text_chunks(text, 24_000)
        endpoint = (
            "https://api-free.deepl.com/v2/translate"
            if use_free_api
            else "https://api.deepl.com/v2/translate"
        )
        translated: list[str] = []
        for index, chunk in enumerate(chunks):
            check_cancelled(cancel)
            payload = json.dumps(
                {"text": [chunk], "target_lang": target_language.upper()}, ensure_ascii=False
            ).encode("utf-8")
            request = urllib.request.Request(
                endpoint,
                data=payload,
                headers={
                    "Authorization": f"DeepL-Auth-Key {api_key.strip()}",
                    "Content-Type": "application/json",
                    "User-Agent": "AIO-Media-Tool/0.3",
                },
                method="POST",
            )
            try:
                with urllib.request.urlopen(request, timeout=45) as response:
                    result = json.loads(response.read().decode("utf-8"))
            except urllib.error.HTTPError as exc:
                detail = exc.read().decode("utf-8", errors="replace")[:800]
                raise RuntimeError(
                    f"DeepL hat die Anfrage abgelehnt ({exc.code}): {detail}"
                ) from exc
            except (OSError, urllib.error.URLError, json.JSONDecodeError) as exc:
                raise RuntimeError(f"DeepL-Übersetzung fehlgeschlagen: {exc}") from exc
            try:
                translated.append(str(result["translations"][0]["text"]))
            except (KeyError, IndexError, TypeError) as exc:
                raise RuntimeError("DeepL hat eine unerwartete Antwort geliefert.") from exc
            progress(int((index + 1) / len(chunks) * 100), "DeepL übersetzt Text")
        return "".join(translated)

    @staticmethod
    def translate_marian(
        text: str,
        model_path: Path,
        progress: ProgressCallback = noop_progress,
        cancel: Event | None = None,
    ) -> str:
        if not model_path.is_dir():
            raise FileNotFoundError(
                "Bitte einen bereits lokal gespeicherten MarianMT-Modellordner wählen."
            )
        try:
            from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
        except (ImportError, OSError) as exc:
            raise RuntimeError(
                "Für MarianMT fehlen lokale Übersetzungsabhängigkeiten "
                "('uv sync --extra local-translation' plus eine zur Hardware passende PyTorch-Installation)."
            ) from exc
        tokenizer = AutoTokenizer.from_pretrained(str(model_path), local_files_only=True)
        model = AutoModelForSeq2SeqLM.from_pretrained(str(model_path), local_files_only=True)
        chunks = OCRService._text_chunks(text, 1600)
        output: list[str] = []
        for index, chunk in enumerate(chunks):
            check_cancelled(cancel)
            encoded = tokenizer(chunk, return_tensors="pt", truncation=True, max_length=512)
            generated = model.generate(**encoded, max_new_tokens=512)
            output.append(tokenizer.decode(generated[0], skip_special_tokens=True))
            progress(int((index + 1) / len(chunks) * 100), "MarianMT übersetzt lokal")
        return "\n".join(output)

    @staticmethod
    def _text_chunks(text: str, maximum: int) -> list[str]:
        chunks: list[str] = []
        remaining = text
        while len(remaining) > maximum:
            split = remaining.rfind("\n", 0, maximum)
            if split < maximum // 2:
                split = remaining.rfind(". ", 0, maximum)
                split = split + 2 if split >= 0 else maximum
            chunks.append(remaining[:split])
            remaining = remaining[split:]
        if remaining:
            chunks.append(remaining)
        return chunks or [""]

    @staticmethod
    def export_txt(original: str, translation: str, output: Path) -> Path:
        if translation.strip():
            text = f"ORIGINAL\n========\n{original}\n\nÜBERSETZUNG\n===========\n{translation}\n"
        else:
            text = original.rstrip() + "\n"
        return atomic_write_bytes(output.with_suffix(".txt"), text.encode("utf-8"))

    @staticmethod
    def export_docx(original: str, translation: str, output: Path) -> Path:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx ist nicht installiert.") from exc
        output = unique_output(output.with_suffix(".docx"))
        output.parent.mkdir(parents=True, exist_ok=True)
        handle, temp_name = tempfile.mkstemp(
            prefix=f".{output.stem}-", suffix=".docx", dir=output.parent
        )
        os.close(handle)
        temp = Path(temp_name)
        try:
            document = Document()
            document.add_heading("OCR – Original", level=1)
            for paragraph in original.split("\n"):
                document.add_paragraph(paragraph)
            if translation.strip():
                document.add_page_break()
                document.add_heading("Übersetzung", level=1)
                for paragraph in translation.split("\n"):
                    document.add_paragraph(paragraph)
            document.save(temp)
            os.replace(temp, output)
        finally:
            temp.unlink(missing_ok=True)
        return output
